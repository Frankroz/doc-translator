import os
import threading
import time
import copy
import openpyxl
from openpyxl.styles import Alignment
from openpyxl.cell.cell import MergedCell
from docx import Document
from deep_translator import GoogleTranslator
import customtkinter as ctk
from tkinter import filedialog, messagebox

# Intentar cargar win32com para soporte nativo de Office (obligatorio para .doc/.xls)
try:
    import win32com.client as win32
except ImportError:
    win32 = None

# Configuración Visual Premium
ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("blue")

class DocTranslatorPro(ctk.CTk):
    def __init__(self):
        super().__init__()

        self.title("With Love ❤︎ - Doc Translator")
        self.geometry("1000x800")
        
        self.is_cancelled = False
        self.translator = GoogleTranslator(source='auto', target='es')
        self.files_path_set = set()
        self.files_to_process = []
        
        self.total_units = 0
        self.done_units = 0
        self.start_time = 0

        self._build_ui()

    def _build_ui(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(2, weight=1)

        # CABECERA
        self.head = ctk.CTkLabel(self, text="DOC TRANSLATOR AI", font=ctk.CTkFont(size=32, weight="bold"))
        self.head.grid(row=0, column=0, pady=(40, 20))

        # ZONA DE CARGA
        self.up_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.up_frame.grid(row=1, column=0, padx=20, pady=10)

        self.btn_fold = ctk.CTkButton(self.up_frame, text="📂 Carpeta", command=self.add_folder, fg_color="#27ae60", width=150)
        self.btn_fold.grid(row=0, column=0, padx=10)

        self.btn_one = ctk.CTkButton(self.up_frame, text="📄 Archivo", command=self.add_file, fg_color="#2980b9", width=150)
        self.btn_one.grid(row=0, column=1, padx=10)

        self.btn_clr = ctk.CTkButton(self.up_frame, text="🗑️ Limpiar", command=self.clear, fg_color="#7f8c8d", width=150)
        self.btn_clr.grid(row=0, column=2, padx=10)

        # LISTA DE ARCHIVOS
        self.list_frame = ctk.CTkScrollableFrame(self, label_text="Cola de Documentos", label_font=ctk.CTkFont(weight="bold"))
        self.list_frame.grid(row=2, column=0, padx=50, pady=20, sticky="nsew")

        # PANEL DE CONTROL Y PROGRESO
        self.bot_panel = ctk.CTkFrame(self, corner_radius=20, fg_color="#1a1a1a")
        self.bot_panel.grid(row=3, column=0, padx=40, pady=30, sticky="ew")
        self.bot_panel.grid_columnconfigure(0, weight=1)

        self.lbl_status = ctk.CTkLabel(self.bot_panel, text="Listo para procesar", font=ctk.CTkFont(size=16))
        self.lbl_status.grid(row=0, column=0, pady=(20, 5))

        self.lbl_eta = ctk.CTkLabel(self.bot_panel, text="ETA: --:--", text_color="#95a5a6")
        self.lbl_eta.grid(row=1, column=0, pady=(0, 10))

        self.bar = ctk.CTkProgressBar(self.bot_panel, width=800, height=12)
        self.bar.set(0)
        self.bar.grid(row=2, column=0, pady=10, padx=30)

        self.ctrl_btns = ctk.CTkFrame(self.bot_panel, fg_color="transparent")
        self.ctrl_btns.grid(row=3, column=0, pady=(10, 25))

        self.btn_go = ctk.CTkButton(self.ctrl_btns, text="🚀 EMPEZAR TRADUCCIÓN GLOBAL", command=self.start, 
                                    width=300, height=50, font=ctk.CTkFont(size=16, weight="bold"), state="disabled")
        self.btn_go.grid(row=0, column=0, padx=10)

        self.btn_stop = ctk.CTkButton(self.ctrl_btns, text="🛑 CANCELAR", command=self.stop, 
                                      width=150, height=50, fg_color="#c0392b", state="disabled")
        self.btn_stop.grid(row=0, column=1, padx=10)

    # --- LÓGICA DE GESTIÓN DE ARCHIVOS ---
    def add_folder(self):
        f = filedialog.askdirectory()
        if f:
            for file in os.listdir(f):
                if file.lower().endswith(('.docx', '.doc', '.xlsx', '.xls')) and not file.startswith(('~$', 'BILINGUAL_')):
                    self.register_file(os.path.join(f, file))

    def add_file(self):
        f = filedialog.askopenfilename(filetypes=[("Documentos Office", "*.docx *.doc *.xlsx *.xls")])
        if f and not os.path.basename(f).startswith('BILINGUAL_'):
            self.register_file(f)

    def register_file(self, path):
        abs_p = os.path.abspath(path)
        if abs_p not in self.files_path_set:
            self.files_path_set.add(abs_p)
            v = ctk.BooleanVar(value=True)
            chk = ctk.CTkCheckBox(self.list_frame, text=os.path.basename(path), variable=v)
            chk.pack(fill="x", padx=15, pady=5)
            self.files_to_process.append((abs_p, v))
            self.btn_go.configure(state="normal")

    def clear(self):
        for w in self.list_frame.winfo_children(): w.destroy()
        self.files_to_process, self.files_path_set = [], set()
        self.btn_go.configure(state="disabled")

    def stop(self):
        self.is_cancelled = True
        self.lbl_status.configure(text="Cancelando...", text_color="#e74c3c")

    # --- NÚCLEO DE TRADUCCIÓN ---
    def start(self):
        sel = [p for p, v in self.files_to_process if v.get()]
        if not sel: return
        self.is_cancelled = False
        self.btn_go.configure(state="disabled"); self.btn_stop.configure(state="normal")
        threading.Thread(target=self.main_loop, args=(sel,), daemon=True).start()

    def main_loop(self, files):
        self.lbl_status.configure(text="Calculando carga de trabajo...")
        self.total_units = 0
        self.done_units = 0
        
        ready_queue = []
        # FASE 1: Conversión y Análisis
        for p in files:
            if self.is_cancelled: break
            units, temp_p = self.prepare_document(p)
            self.total_units += units
            ready_queue.append((p, temp_p))

        if self.total_units == 0: self.total_units = 1
        self.start_time = time.time()

        # FASE 2: Traducción
        for orig, temp in ready_queue:
            if self.is_cancelled: break
            self.lbl_status.configure(text=f"Procesando: {os.path.basename(orig)}")
            try:
                if temp.endswith('.docx'):
                    self.translate_word(orig, temp)
                else:
                    self.translate_excel(orig, temp)
            except Exception as e:
                print(f"Error: {e}")

        # Limpieza de UI
        self.bar.set(1 if not self.is_cancelled else 0)
        self.lbl_status.configure(text="¡Tarea Finalizada!" if not self.is_cancelled else "Proceso Interrumpido", 
                                   text_color="#2ecc71" if not self.is_cancelled else "#e74c3c")
        self.lbl_eta.configure(text="100% Completado")
        self.btn_go.configure(state="normal"); self.btn_stop.configure(state="disabled")
        messagebox.showinfo("DocTranslator", "El proceso global ha terminado.")

    def prepare_document(self, path):
        """Crea documento temporal si es legado y cuenta unidades de trabajo."""
        work_path = path
        # Manejo de .doc y .xls antiguos
        if path.lower().endswith(('.doc', '.xls')):
            if win32:
                try:
                    if path.lower().endswith('.doc'):
                        app = win32.gencache.EnsureDispatch('Word.Application')
                        d = app.Documents.Open(os.path.abspath(path))
                        work_path = os.path.abspath(path) + "_TEMP_.docx"
                        d.SaveAs2(work_path, FileFormat=16)
                        d.Close(); app.Quit()
                    else:
                        app = win32.gencache.EnsureDispatch('Excel.Application')
                        b = app.Workbooks.Open(os.path.abspath(path))
                        work_path = os.path.abspath(path) + "_TEMP_.xlsx"
                        b.SaveAs(work_path, FileFormat=51)
                        b.Close(); app.Quit()
                except Exception as e:
                    print(f"Error conversión: {e}")
            else:
                print("win32com no disponible")

        # Conteo de unidades
        cnt = 0
        try:
            if work_path.lower().endswith('.docx'):
                d = Document(work_path)
                cnt = len([p for p in d.paragraphs if p.text.strip()])
                for t in d.tables:
                    for r in t.rows:
                        cnt += sum(1 for c in r.cells if c.text.strip())
            else:
                wb = openpyxl.load_workbook(work_path, data_only=True)
                for s in wb.worksheets:
                    for row in s.iter_rows(values_only=True):
                        cnt += sum(1 for c in row if isinstance(c, str) and c.strip())
        except: cnt = 1
        return cnt, work_path

    def translate_word(self, orig, temp):
        doc = Document(temp)
        
        # 1. Traducir Párrafos (Fuera de tablas)
        for p in doc.paragraphs:
            if self.is_cancelled: return
            if p.text.strip():
                try:
                    # Guardamos el estilo del primer run para replicarlo
                    original_style = p.runs[0].font if p.runs else None
                    res = self.translator.translate(p.text)
                    
                    run = p.add_run(f"\n{res}")
                    run.italic = False
                    if original_style:
                        run.font.name = original_style.name
                        run.font.size = original_style.size
                except: pass
                self.upd_prog()

        # 2. Traducir Tablas (Manejo de Merged Cells y Estilos)
        for table in doc.tables:
            # Usamos set para rastrear celdas ya procesadas (evita duplicados en celdas combinadas)
            processed_cells = set()
            
            for row in table.rows:
                for cell in row.cells:
                    if self.is_cancelled: return
                    
                    # El ID de la celda en memoria nos dice si es la misma (merged)
                    cell_id = cell._tc
                    if cell_id in processed_cells:
                        continue
                    
                    if cell.text.strip():
                        try:
                            # 1. Capturamos el texto original
                            original_text = cell.text
                            
                            # 2. Capturamos el formato del primer párrafo de la celda
                            # (Normalmente las celdas tienen al menos un párrafo)
                            first_para = cell.paragraphs[0]
                            
                            # 3. Traducimos
                            res = self.translator.translate(original_text)
                            
                            # 4. Insertamos la traducción manteniendo el estilo
                            # Añadimos un nuevo párrafo o un run al final del existente
                            new_run = first_para.add_run(f"\n{res}")
                            new_run.italic = False
                            
                        except Exception as e:
                            print(f"Error en celda: {e}")
                        
                        self.upd_prog()
                    
                    # Marcamos como procesada
                    processed_cells.add(cell_id)

        # 3. Guardar y Limpiar
        out_name = f"BILINGUAL_{os.path.basename(orig)}"
        if orig.lower().endswith('.doc'): 
            out_name += "x"
            
        output_path = os.path.join(os.path.dirname(orig), out_name)
        doc.save(output_path)
        
        if "_TEMP_" in temp: 
            try: os.remove(temp)
            except: pass

    def translate_excel(self, orig, temp):
        wb = openpyxl.load_workbook(temp)
        new_wb = openpyxl.Workbook()
        
        for idx, sn in enumerate(wb.sheetnames):
            if self.is_cancelled: break
            ws = wb[sn]
            nws = new_wb.active if idx == 0 else new_wb.create_sheet(sn)
            nws.title = sn

            # COPIA DE DIMENSIONES MEJORADA
            for c, d in ws.column_dimensions.items(): nws.column_dimensions[c].width = d.width
            for r, d in ws.row_dimensions.items(): nws.row_dimensions[r].height = d.height

            # DETECCIÓN DE RANGO REAL (Evita cortes a la mitad)
            max_row = ws.max_row
            max_col = ws.max_column
            
            for r in range(1, max_row + 1):
                for c in range(1, max_col + 1):
                    if self.is_cancelled: return
                    cell, ncell = ws.cell(row=r, column=c), nws.cell(row=r, column=c)
                    
                    if not isinstance(cell, MergedCell):
                        if isinstance(cell.value, str) and cell.value.strip():
                            try:
                                res = self.translator.translate(cell.value)
                                ncell.value = f"{cell.value}\n{res}"
                            except: ncell.value = cell.value
                            self.upd_prog()
                        else: ncell.value = cell.value

                    if cell.has_style:
                        ncell.font = copy.copy(cell.font)
                        ncell.border = copy.copy(cell.border)
                        ncell.fill = copy.copy(cell.fill)
                        ncell.alignment = Alignment(
                                wrap_text=True, 
                                horizontal=cell.alignment.horizontal or 'center', 
                                vertical=cell.alignment.vertical or 'center'
                            )

            for mr in ws.merged_cells.ranges: nws.merge_cells(str(mr))
            if hasattr(ws, '_images'):
                    for img in ws._images:
                        ws.add_image(copy.copy(img))

        out_name = f"BILINGUAL_{os.path.basename(orig)}"
        if orig.lower().endswith('.xls'): out_name += "x"
        new_wb.save(os.path.join(os.path.dirname(orig), out_name))
        if "_TEMP_" in temp: os.remove(temp)

    def upd_prog(self):
        self.done_units += 1
        p = self.done_units / self.total_units
        self.bar.set(p)
        
        elapsed = time.time() - self.start_time
        if self.done_units > 0:
            total_est = (elapsed / self.done_units) * self.total_units
            rem = total_est - elapsed
            m, s = divmod(int(rem), 60)
            self.lbl_eta.configure(text=f"ETA Global: {m:02d}:{s:02d}")
        self.update_idletasks()

if __name__ == "__main__":
    app = DocTranslatorPro()
    app.mainloop()